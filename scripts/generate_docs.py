import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_section(doc, title, content):
    """Add a section to the document with proper formatting."""
    # Add title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content
    paragraph = doc.add_paragraph(content)
    paragraph.style.font.size = Pt(11)

def generate_documentation():
    """Generate a Word document containing all documentation."""
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('DecentraLearn Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('Comprehensive Guide to DecentraLearn Framework')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add table of contents
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. Development Setup Guide')
    doc.add_paragraph('2. Code Style Guide')
    doc.add_paragraph('3. Testing Guide')
    doc.add_paragraph('4. API Documentation')
    doc.add_paragraph('5. Architecture Documentation')
    doc.add_paragraph('6. Tutorials')
    doc.add_paragraph('7. Examples')
    
    # Add page break
    doc.add_page_break()
    
    # Get the workspace root directory
    workspace_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # Read and add development setup guide
    setup_path = os.path.join(workspace_root, 'docs/development/setup.md')
    with open(setup_path, 'r') as f:
        setup_content = f.read()
    add_section(doc, 'Development Setup Guide', setup_content)
    doc.add_page_break()
    
    # Read and add code style guide
    code_style_path = os.path.join(workspace_root, 'docs/development/code_style.md')
    with open(code_style_path, 'r') as f:
        code_style_content = f.read()
    add_section(doc, 'Code Style Guide', code_style_content)
    doc.add_page_break()
    
    # Read and add testing guide
    testing_path = os.path.join(workspace_root, 'docs/development/testing.md')
    with open(testing_path, 'r') as f:
        testing_content = f.read()
    add_section(doc, 'Testing Guide', testing_content)
    doc.add_page_break()
    
    # Read and add API documentation
    api_path = os.path.join(workspace_root, 'docs/api/README.md')
    with open(api_path, 'r') as f:
        api_content = f.read()
    add_section(doc, 'API Documentation', api_content)
    doc.add_page_break()
    
    # Read and add architecture documentation
    arch_path = os.path.join(workspace_root, 'docs/architecture/overview.md')
    with open(arch_path, 'r') as f:
        arch_content = f.read()
    add_section(doc, 'Architecture Documentation', arch_content)
    doc.add_page_break()
    
    # Read and add tutorials
    tutorials_path = os.path.join(workspace_root, 'docs/tutorials/README.md')
    with open(tutorials_path, 'r') as f:
        tutorials_content = f.read()
    add_section(doc, 'Tutorials', tutorials_content)
    doc.add_page_break()
    
    # Read and add examples
    examples_path = os.path.join(workspace_root, 'docs/examples/README.md')
    with open(examples_path, 'r') as f:
        examples_content = f.read()
    add_section(doc, 'Examples', examples_content)
    
    # Save the document
    output_path = os.path.join(workspace_root, 'DecentraLearn_Documentation.docx')
    doc.save(output_path)
    print(f"Documentation generated successfully at: {output_path}")

if __name__ == '__main__':
    generate_documentation()